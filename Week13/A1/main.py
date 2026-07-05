"""
LLM-Powered CV Feedback and Optimization Backend
YCCIA 2511 MSE PSE - CV Feedback Activity

Purpose:
- Read a user's CV and a target job description from PDF, DOCX, or TXT.
- Calculate richer ATS and CV structure scores.
- Generate constructive, personalized CV feedback using Gemini or OpenAI when an API key is available.
- Use the generated feedback to rewrite and optimize the user's ACTUAL CV.
- Export feedback, JSON analysis, optimized DOCX, optimized PDF, and logs.

Environment variables:
- LLM_PROVIDER=gemini or openai
- GEMINI_API_KEY=your_key_here
- GEMINI_MODEL=gemini-2.0-flash
- OPENAI_API_KEY=your_key_here
- OPENAI_MODEL=gpt-4.1-mini

Example:
python main.py --cv resume.pdf --jd job_description.pdf --provider gemini --output output --docx --pdf --markdown --json
"""

from __future__ import annotations

import argparse
import json
import logging
import os
import re
import time
from collections import Counter
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path
from textwrap import wrap
from typing import Final, List, Optional, Protocol

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


# =====================================================
# 0) CONFIGURATION
# =====================================================

DEFAULT_LLM_PROVIDER: Final[str] = os.getenv("LLM_PROVIDER", "gemini").strip().lower()
GEMINI_API_KEY: Final[str] = os.getenv("GEMINI_API_KEY", "").strip()
GEMINI_MODEL: Final[str] = os.getenv("GEMINI_MODEL", "gemini-2.0-flash").strip()
OPENAI_API_KEY: Final[str] = os.getenv("OPENAI_API_KEY", "").strip()
OPENAI_MODEL: Final[str] = os.getenv("OPENAI_MODEL", "gpt-4.1-mini").strip()
SUPPORTED_EXTS: Final[set[str]] = {".pdf", ".docx", ".txt"}
DEFAULT_OUTPUT_DIR: Final[Path] = Path("output")

STOPWORDS: Final[frozenset[str]] = frozenset(
    {
        "and", "or", "the", "a", "an", "to", "for", "of", "in", "on", "with", "as", "at", "by",
        "is", "are", "be", "this", "that", "it", "we", "you", "your", "our", "will", "from", "using",
        "work", "team", "teams", "role", "job", "title", "location", "type", "summary", "skills", "required",
        "preferred", "education", "experience", "knowledge", "understanding", "familiarity", "degree", "field",
        "candidate", "applicant", "responsibilities", "requirements", "including", "ability", "strong", "good",
        "excellent", "within", "across", "based", "must", "should", "have", "has", "their", "they", "them",
    }
)
TOKEN_PATTERN: Final[re.Pattern[str]] = re.compile(r"[A-Za-z][A-Za-z\+\#\.\-]{1,}")
EMAIL_PHONE_PATTERN: Final[re.Pattern[str]] = re.compile(r"[\w\.-]+@[\w\.-]+|\+?\d[\d\s\-\(\)]{7,}", re.I)

TECH_KEYWORDS: Final[set[str]] = {
    "python", "java", "c#", "c++", "javascript", "typescript", "sql", "html", "css", "react", "angular",
    "node", "django", "flask", "fastapi", "spring", "api", "rest", "restful", "microservices", "backend",
    "frontend", "full-stack", "git", "github", "gitlab", "docker", "kubernetes", "ci/cd", "devops", "aws", "azure",
    "gcp", "cloud", "linux", "windows", "testing", "unit", "integration", "pytest", "selenium", "agile", "scrum",
    "oop", "object-oriented", "database", "sql server", "mysql", "postgresql", "sqlite", "mongodb", "redis", "rabbitmq",
    "etl", "pandas", "numpy", "scikit-learn", "machine learning", "data", "analytics", "security", "performance",
}
SOFT_KEYWORDS: Final[set[str]] = {
    "communication", "collaboration", "stakeholder", "teamwork", "problem-solving", "documentation", "leadership",
    "support", "mentoring", "ownership", "analysis", "planning", "delivery", "customer", "business",
}
ACTION_VERBS: Final[set[str]] = {
    "developed", "designed", "built", "created", "implemented", "maintained", "improved", "automated", "optimized",
    "tested", "deployed", "managed", "collaborated", "documented", "integrated", "migrated", "supported", "resolved",
}


# =====================================================
# 1) LOGGING
# =====================================================

def setup_logging(output_dir: Path) -> Path:
    """Configure application logging."""
    log_dir = output_dir / "logs"
    log_dir.mkdir(parents=True, exist_ok=True)
    log_path = log_dir / "analysis.log"
    logging.basicConfig(
        filename=str(log_path),
        level=logging.INFO,
        format="%(asctime)s | %(levelname)s | %(message)s",
        force=True,
    )
    return log_path


# =====================================================
# 2) TEXT EXTRACTION
# =====================================================

class TextExtractor:
    """Base extractor interface."""

    def extract(self, file_path: Path) -> str:
        """Read text from a file."""
        raise NotImplementedError


class PDFTextExtractor(TextExtractor):
    """Extract text from PDF files."""

    def extract(self, file_path: Path) -> str:
        try:
            import pdfplumber  # type: ignore
        except ImportError as exc:
            raise ImportError("Missing dependency: pdfplumber. Install with: pip install pdfplumber") from exc

        parts: list[str] = []
        with pdfplumber.open(str(file_path)) as pdf:
            for page in pdf.pages:
                parts.append((page.extract_text() or "").strip())
        return "\n".join(part for part in parts if part).strip()


class DOCXTextExtractor(TextExtractor):
    """Extract text from DOCX files."""

    def extract(self, file_path: Path) -> str:
        document = Document(str(file_path))
        return "\n".join(p.text for p in document.paragraphs if p.text.strip()).strip()


class TXTTextExtractor(TextExtractor):
    """Extract text from TXT files."""

    def extract(self, file_path: Path) -> str:
        return file_path.read_text(encoding="utf-8", errors="ignore").strip()


class ExtractorFactory:
    """Factory for file text extraction."""

    @staticmethod
    def get(file_path: Path) -> TextExtractor:
        ext = file_path.suffix.lower()
        if ext == ".pdf":
            return PDFTextExtractor()
        if ext == ".docx":
            return DOCXTextExtractor()
        if ext == ".txt":
            return TXTTextExtractor()
        raise ValueError(f"Unsupported file format: {ext}. Use PDF, DOCX, or TXT.")


def read_document(path_str: str) -> str:
    """Read a supported document path and return clean text."""
    file_path = Path(path_str.strip().strip('"')).expanduser()
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    if file_path.suffix.lower() not in SUPPORTED_EXTS:
        raise ValueError(f"Unsupported file type: {file_path.suffix}. Use PDF, DOCX, or TXT.")
    text = ExtractorFactory.get(file_path).extract(file_path)
    if not text:
        raise ValueError(f"No readable text found in file: {file_path}")
    return text


# =====================================================
# 3) ATS SCORING AND CV STRUCTURE ANALYSIS
# =====================================================

def normalize_term(term: str) -> str:
    """Normalize keywords for matching."""
    return term.lower().strip().replace("ci cd", "ci/cd").replace("rest api", "rest")


def tokenize(text: str) -> List[str]:
    """Convert raw text to ATS-style keyword tokens."""
    lower_text = text.lower().replace("c sharp", "c#").replace("dot net", ".net")
    words = TOKEN_PATTERN.findall(lower_text)
    clean_words = [normalize_term(w) for w in words if w not in STOPWORDS and len(w) >= 2]

    # Add common phrase tokens because ATS and recruiters often search phrases, not only single words.
    phrase_candidates = [
        "sql server", "rest api", "restful api", "machine learning", "data analytics", "object oriented",
        "object-oriented", "unit testing", "integration testing", "ci/cd", "full stack", "full-stack",
        "software engineering", "cloud computing", "business analysis", "technical documentation",
    ]
    for phrase in phrase_candidates:
        if phrase in lower_text:
            clean_words.append(normalize_term(phrase))
    return clean_words


@dataclass(frozen=True, slots=True)
class ATSResult:
    """Rich ATS scoring result."""

    overall_score: int
    keyword_score: int
    skills_score: int
    experience_score: int
    formatting_score: int
    readability_score: int
    section_score: int
    matched_keywords: List[str]
    missing_keywords: List[str]
    top_keywords: List[str]
    matched_technical_keywords: List[str]
    missing_technical_keywords: List[str]


class ATSScorer:
    """ATS scoring using keyword overlap, technical skills, structure, and readability."""

    def score(self, cv_text: str, jd_text: str, profile: "CVProfile", top_n: int = 45) -> ATSResult:
        if top_n <= 0:
            raise ValueError("top_n must be greater than 0")

        cv_tokens = set(tokenize(cv_text))
        jd_tokens = tokenize(jd_text)
        if not jd_tokens:
            return ATSResult(0, 0, 0, 0, 0, 0, 0, [], [], [], [], [])

        top_keywords = [word for word, _ in Counter(jd_tokens).most_common(top_n)]
        matched = [keyword for keyword in top_keywords if keyword in cv_tokens]
        missing = [keyword for keyword in top_keywords if keyword not in cv_tokens]
        keyword_score = round((len(matched) / len(top_keywords)) * 100) if top_keywords else 0

        jd_tech = sorted({kw for kw in TECH_KEYWORDS if kw in jd_text.lower() or kw in jd_tokens})
        if not jd_tech:
            jd_tech = sorted({kw for kw in TECH_KEYWORDS if kw in top_keywords})
        matched_tech = [kw for kw in jd_tech if kw in cv_text.lower() or kw in cv_tokens]
        missing_tech = [kw for kw in jd_tech if kw not in matched_tech]
        skills_score = round((len(matched_tech) / len(jd_tech)) * 100) if jd_tech else keyword_score

        action_hits = sum(1 for verb in ACTION_VERBS if re.search(rf"\b{re.escape(verb)}\b", cv_text, re.I))
        experience_score = min(100, round((profile.bullet_count * 8) + (action_hits * 5)))

        formatting_score = self._formatting_score(cv_text, profile)
        readability_score = self._readability_score(cv_text)
        section_score = profile.section_score
        overall = round(
            keyword_score * 0.35
            + skills_score * 0.20
            + experience_score * 0.15
            + formatting_score * 0.10
            + readability_score * 0.10
            + section_score * 0.10
        )

        return ATSResult(
            overall_score=overall,
            keyword_score=keyword_score,
            skills_score=skills_score,
            experience_score=experience_score,
            formatting_score=formatting_score,
            readability_score=readability_score,
            section_score=section_score,
            matched_keywords=matched,
            missing_keywords=missing,
            top_keywords=top_keywords,
            matched_technical_keywords=matched_tech,
            missing_technical_keywords=missing_tech,
        )

    @staticmethod
    def _formatting_score(cv_text: str, profile: "CVProfile") -> int:
        score = 100
        if "│" in cv_text or "|" in cv_text and cv_text.count("|") > 15:
            score -= 15
        if profile.bullet_count < 5:
            score -= 15
        if profile.word_count < 180:
            score -= 20
        if profile.word_count > 1200:
            score -= 10
        if not profile.has_contact:
            score -= 15
        return max(0, min(100, score))

    @staticmethod
    def _readability_score(cv_text: str) -> int:
        sentences = max(1, len(re.findall(r"[.!?]", cv_text)))
        words = max(1, len(tokenize(cv_text)))
        avg_sentence_len = words / sentences
        if 8 <= avg_sentence_len <= 22:
            return 95
        if 23 <= avg_sentence_len <= 30:
            return 80
        if avg_sentence_len > 30:
            return 65
        return 75


@dataclass(frozen=True, slots=True)
class CVProfile:
    """Structured CV analysis used for feedback."""

    has_summary: bool
    has_skills: bool
    has_experience: bool
    has_projects: bool
    has_education: bool
    has_certifications: bool
    has_achievements: bool
    has_contact: bool
    has_linkedin: bool
    has_github: bool
    has_portfolio: bool
    bullet_count: int
    word_count: int
    section_score: int
    missing_sections: List[str]


class CVStructureAnalyzer:
    """Check common CV structure and presentation issues."""

    SECTION_PATTERNS: Final[dict[str, re.Pattern[str]]] = {
        "summary": re.compile(r"\b(summary|profile|professional summary|summary of qualifications)\b", re.I),
        "skills": re.compile(r"\b(skills|technical skills|computer skills|core skills)\b", re.I),
        "experience": re.compile(r"\b(experience|employment|work history|professional experience|career history)\b", re.I),
        "projects": re.compile(r"\b(projects|selected projects|academic projects|portfolio projects)\b", re.I),
        "education": re.compile(r"\b(education|qualification|qualifications|academic background)\b", re.I),
        "certifications": re.compile(r"\b(certification|certifications|certificate|certificates|training)\b", re.I),
        "achievements": re.compile(r"\b(achievement|achievements|awards|accomplishments|highlights)\b", re.I),
        "contact": EMAIL_PHONE_PATTERN,
        "linkedin": re.compile(r"linkedin\.com|\blinkedin\b", re.I),
        "github": re.compile(r"github\.com|\bgithub\b", re.I),
        "portfolio": re.compile(r"portfolio|personal website|github\.io|https?://", re.I),
    }

    def analyse(self, cv_text: str) -> CVProfile:
        bullet_count = len(re.findall(r"(^|\n)\s*(•|-|\*)\s+", cv_text))
        word_count = len(tokenize(cv_text))
        flags = {key: bool(pattern.search(cv_text)) for key, pattern in self.SECTION_PATTERNS.items()}

        important_sections = ["summary", "skills", "experience", "education", "contact"]
        optional_sections = ["projects", "certifications", "achievements", "linkedin", "github", "portfolio"]
        section_score = round(
            (sum(1 for key in important_sections if flags[key]) / len(important_sections)) * 75
            + (sum(1 for key in optional_sections if flags[key]) / len(optional_sections)) * 25
        )
        missing_sections = [key.replace("_", " ").title() for key in important_sections + optional_sections if not flags[key]]

        return CVProfile(
            has_summary=flags["summary"],
            has_skills=flags["skills"],
            has_experience=flags["experience"],
            has_projects=flags["projects"],
            has_education=flags["education"],
            has_certifications=flags["certifications"],
            has_achievements=flags["achievements"],
            has_contact=flags["contact"],
            has_linkedin=flags["linkedin"],
            has_github=flags["github"],
            has_portfolio=flags["portfolio"],
            bullet_count=bullet_count,
            word_count=word_count,
            section_score=section_score,
            missing_sections=missing_sections,
        )


# =====================================================
# 4) LLM CLIENTS
# =====================================================

class LLMClient(Protocol):
    """LLM protocol."""

    def generate(self, prompt: str) -> str:
        """Return generated text."""
        raise NotImplementedError


class GeminiClient:
    """Google Gemini client."""

    def __init__(self, api_key: str, model: str) -> None:
        if not api_key:
            raise ValueError("GEMINI_API_KEY is missing")
        from google import genai  # type: ignore
        self._client = genai.Client(api_key=api_key)
        self._model = model

    def generate(self, prompt: str) -> str:
        response = self._client.models.generate_content(model=self._model, contents=prompt)
        return (getattr(response, "text", "") or "").strip()


class OpenAIClient:
    """OpenAI client."""

    def __init__(self, api_key: str, model: str) -> None:
        if not api_key:
            raise ValueError("OPENAI_API_KEY is missing")
        from openai import OpenAI  # type: ignore
        self._client = OpenAI(api_key=api_key)
        self._model = model

    def generate(self, prompt: str) -> str:
        response = self._client.chat.completions.create(
            model=self._model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a senior technical recruiter, ATS resume specialist, "
                        "and career coach for software engineering roles. You must be truthful and never invent CV facts."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.25,
        )
        return (response.choices[0].message.content or "").strip()


class LLMFactory:
    """Create the selected LLM provider if keys are available."""

    @staticmethod
    def create(provider: str) -> Optional[LLMClient]:
        clean_provider = provider.strip().lower()
        if clean_provider == "openai" and OPENAI_API_KEY:
            return OpenAIClient(OPENAI_API_KEY, OPENAI_MODEL)
        if clean_provider == "gemini" and GEMINI_API_KEY:
            return GeminiClient(GEMINI_API_KEY, GEMINI_MODEL)
        return None


# =====================================================
# 5) PROMPT ENGINEERING
# =====================================================

class PromptBuilder:
    """Build prompts for feedback generation and CV rewriting."""

    @staticmethod
    def feedback_prompt(cv_text: str, jd_text: str, ats: ATSResult, profile: CVProfile) -> str:
        missing = ", ".join(ats.missing_keywords[:30]) or "None"
        matched = ", ".join(ats.matched_keywords[:30]) or "None"
        missing_sections = ", ".join(profile.missing_sections) or "None"
        missing_tech = ", ".join(ats.missing_technical_keywords[:25]) or "None"
        return f"""
ROLE:
You are a senior technical recruiter, ATS resume specialist, and software engineering career coach.

GOAL:
Generate constructive, personalized feedback on the user's CV and provide clear recommendations to improve its content, structure, and presentation for the target role.

IMPORTANT RULES:
1. Be specific and practical. Do not give generic advice only.
2. Do not invent fake employment history, qualifications, companies, dates, projects, metrics, or certifications.
3. You may rewrite wording to be clearer, stronger, and more ATS friendly.
4. If the CV does not show enough evidence for a requirement, say "add evidence if true".
5. Use simple professional English.
6. Focus on ATS keyword alignment, technical relevance, measurable achievements, structure, readability, presentation, and ethics.
7. Keep recommendations truthful and suitable for a real job application.

ATS ANALYSIS:
- Overall ATS Score: {ats.overall_score}%
- Keyword Score: {ats.keyword_score}%
- Skills Score: {ats.skills_score}%
- Experience Score: {ats.experience_score}%
- Formatting Score: {ats.formatting_score}%
- Readability Score: {ats.readability_score}%
- Section Score: {ats.section_score}%
- Matched keywords: {matched}
- Missing keywords: {missing}
- Missing technical keywords: {missing_tech}

CV STRUCTURE CHECK:
- Has summary/profile: {profile.has_summary}
- Has skills section: {profile.has_skills}
- Has experience/work history: {profile.has_experience}
- Has projects: {profile.has_projects}
- Has education: {profile.has_education}
- Has certifications: {profile.has_certifications}
- Has achievements: {profile.has_achievements}
- Has contact details: {profile.has_contact}
- Has LinkedIn: {profile.has_linkedin}
- Has GitHub: {profile.has_github}
- Has portfolio: {profile.has_portfolio}
- Missing sections: {missing_sections}
- Approximate keyword word count: {profile.word_count}
- Bullet count detected: {profile.bullet_count}

OUTPUT FORMAT REQUIRED:
1. Overall CV Assessment
   - 3 to 5 sentences about current CV quality and role fit.
2. ATS Match Feedback
   - Explain the scores and the most important missing keywords.
3. Content Recommendations
   - Summary/profile improvements.
   - Skills improvements.
   - Experience bullet improvements using STAR/action verbs and impact.
   - Project, education, certification, and portfolio improvements.
4. Structure and Presentation Recommendations
   - Section order.
   - Bullet style.
   - Formatting and readability.
   - ATS-friendly layout advice.
5. Rewritten CV Content Suggestions
   - Improved professional summary.
   - Improved technical skills section.
   - 6 to 10 improved bullet examples aligned with the job description.
6. Ethical Accuracy Notes
   - State what must not be invented and where evidence should be added only if true.
7. Final Checklist
   - Concise checklist before submission.

TARGET JOB DESCRIPTION:
{jd_text}

USER CV:
{cv_text}
""".strip()

    @staticmethod
    def rewrite_prompt(cv_text: str, jd_text: str, feedback: str, ats: ATSResult) -> str:
        missing = ", ".join(ats.missing_keywords[:30]) or "None"
        missing_tech = ", ".join(ats.missing_technical_keywords[:25]) or "None"
        return f"""
ROLE:
You are a senior ATS resume writer and technical recruiter.

TASK:
Rewrite and optimize the user's ACTUAL CV using the generated feedback below.

CRITICAL RULES:
1. Rewrite the user's actual CV, not a generic template.
2. Preserve all truthful facts from the original CV.
3. Do NOT invent employers, job titles, dates, education, projects, certifications, achievements, metrics, or tools.
4. Add missing ATS keywords naturally only when they are already supported by the CV or write "Add evidence if true".
5. Improve grammar, professional tone, clarity, section order, bullet quality, and ATS compatibility.
6. Use one-column ATS-friendly structure.
7. No tables, no icons, no images, no text boxes.
8. Use strong action verbs and STAR-style bullets where possible.
9. Quantify achievements only if the original CV includes evidence. Otherwise write "Add measurable result if available".
10. Return ONLY the finished rewritten CV. Do not include explanation, markdown fences, or comments outside the CV.

RECOMMENDED CV SECTION ORDER:
Name and Contact Details
Professional Summary
Technical Skills
Professional Experience
Projects
Education
Certifications
Additional Information

ATS TARGETS:
- Missing general keywords to consider truthfully: {missing}
- Missing technical keywords to consider truthfully: {missing_tech}

GENERATED FEEDBACK TO APPLY:
{feedback}

TARGET JOB DESCRIPTION:
{jd_text}

ORIGINAL USER CV:
{cv_text}
""".strip()


# =====================================================
# 6) FEEDBACK GENERATION AND CV REWRITING
# =====================================================

class FeedbackGenerator:
    """Generate LLM feedback or rule-based feedback."""

    def __init__(self, llm: Optional[LLMClient]) -> None:
        self._llm = llm

    def generate(self, prompt: str, ats: ATSResult, profile: CVProfile) -> str:
        if self._llm:
            generated = self._llm.generate(prompt)
            if generated:
                return generated
        return self._fallback_feedback(ats, profile)

    @staticmethod
    def _fallback_feedback(ats: ATSResult, profile: CVProfile) -> str:
        missing_text = ", ".join(ats.missing_keywords[:18]) or "No major missing keyword found."
        missing_sections = ", ".join(profile.missing_sections[:10]) or "No major missing section found."
        return f"""
1. Overall CV Assessment
The CV has useful information, but it should be better aligned to the target job description. The overall ATS score is {ats.overall_score}%, with keyword score {ats.keyword_score}% and skills score {ats.skills_score}%. The CV should improve keyword alignment, technical evidence, action verbs, measurable results, and ATS-friendly structure.

2. ATS Match Feedback
Important missing or weak keywords: {missing_text}. Add these words naturally only when they are true and supported by your real experience.

3. Content Recommendations
- Add or improve the professional summary so it clearly states the target role, core technical skills, and career value.
- Group technical skills by Programming, Backend/API, Databases, DevOps/Cloud, Testing, Tools, and Methods.
- Rewrite experience bullets using action verbs such as developed, designed, maintained, automated, improved, collaborated, tested, documented, and supported.
- Add measurable impact if available, such as time saved, users supported, reports improved, defects reduced, or performance improved.
- Include projects, GitHub, portfolio, and certifications if they are true and relevant.

4. Structure and Presentation Recommendations
- Missing or weak sections: {missing_sections}.
- Recommended order: Contact Details, Professional Summary, Technical Skills, Professional Experience, Projects, Education, Certifications, Additional Information.
- Use one-column formatting, normal headings, simple bullets, and no tables/images/icons.
- Keep bullets short and focused on result, technology, and impact.

5. Rewritten CV Content Suggestions
- Add a role-targeted summary with exact truthful skills.
- Improve skills section by matching the job description keywords.
- Convert task-based bullets into achievement-based bullets.
- Add "Add evidence if true" where the job asks for skills not clearly proven in the current CV.

6. Ethical Accuracy Notes
Do not invent projects, companies, dates, certifications, or numbers. Only add tools and achievements that are true.

7. Final Checklist
- Check contact details.
- Add missing ATS keywords truthfully.
- Improve bullets with action verbs.
- Add measurable impact if available.
- Export as DOCX and PDF.
""".strip()


class CVRewriter:
    """Rewrite the original CV using feedback. Uses LLM when available and fallback when unavailable."""

    def __init__(self, llm: Optional[LLMClient]) -> None:
        self._llm = llm

    def rewrite(self, rewrite_prompt: str, original_cv: str, feedback: str, ats: ATSResult, profile: CVProfile) -> str:
        if self._llm:
            optimized = self._llm.generate(rewrite_prompt).strip()
            if optimized:
                return self._clean_optimized_text(optimized)
        return self._fallback_rewrite(original_cv, feedback, ats, profile)

    @staticmethod
    def _clean_optimized_text(text: str) -> str:
        text = re.sub(r"^```(?:text|markdown)?", "", text.strip(), flags=re.I)
        text = re.sub(r"```$", "", text.strip())
        return text.strip()

    @staticmethod
    def _fallback_rewrite(original_cv: str, feedback: str, ats: ATSResult, profile: CVProfile) -> str:
        """Fallback creates a truthful optimized CV draft from original text and analysis, without inventing facts."""
        lines = [line.strip() for line in original_cv.splitlines() if line.strip()]
        contact_lines = [line for line in lines[:8] if EMAIL_PHONE_PATTERN.search(line) or "linkedin" in line.lower() or "github" in line.lower()]
        possible_name = lines[0] if lines else "Your Name"
        if EMAIL_PHONE_PATTERN.search(possible_name):
            possible_name = "Your Name"

        top_missing = ", ".join(ats.missing_keywords[:10]) or "target job keywords"
        summary = (
            "Professional with software engineering, data, and technical problem-solving experience. "
            "Able to work with structured information, improve processes, document solutions, and support reliable systems. "
            f"Add evidence if true for these target areas: {top_missing}."
        )

        skill_tokens = sorted(set(ats.matched_technical_keywords + [kw for kw in TECH_KEYWORDS if kw in original_cv.lower()]))
        if not skill_tokens:
            skill_tokens = ["Add technical skills from your real experience"]

        bullet_source = [line for line in lines if re.match(r"^(•|-|\*)\s+", line)]
        if not bullet_source:
            bullet_source = lines[1:8]
        improved_bullets = []
        for line in bullet_source[:8]:
            clean = re.sub(r"^(•|-|\*)\s+", "", line).strip()
            if not clean:
                continue
            improved_bullets.append(f"Improved and documented work related to: {clean}. Add measurable result if available.")

        if not improved_bullets:
            improved_bullets = [
                "Add 3 to 6 achievement-focused bullets for each role or project.",
                "Describe the technology used, action completed, and measurable result if available.",
            ]

        contact_text = " | ".join(contact_lines) if contact_lines else "Add email | phone | city | LinkedIn | GitHub/Portfolio if available"
        return f"""
{possible_name.upper()}
{contact_text}

PROFESSIONAL SUMMARY
{summary}

TECHNICAL SKILLS
{', '.join(skill_tokens)}

PROFESSIONAL EXPERIENCE / PROJECTS
{chr(10).join('- ' + bullet for bullet in improved_bullets)}

EDUCATION
Add or keep your real education details from the original CV.

ADDITIONAL INFORMATION
Add certifications, portfolio, GitHub, awards, volunteer work, or languages only if true.

ORIGINAL CV CONTENT TO PRESERVE AND REFINE
{original_cv}
""".strip()


# =====================================================
# 7) EXPORTERS
# =====================================================

class ReportExporter:
    """Export feedback as markdown and text."""

    @staticmethod
    def save(feedback: str, output_dir: Path) -> tuple[Path, Path]:
        output_dir.mkdir(parents=True, exist_ok=True)
        md_path = output_dir / "cv_feedback_report.md"
        txt_path = output_dir / "cv_feedback_report.txt"
        md_path.write_text("# CV Feedback Report\n\n" + feedback + "\n", encoding="utf-8")
        txt_path.write_text(feedback + "\n", encoding="utf-8")
        return md_path, txt_path


class JSONExporter:
    """Export machine-readable analysis."""

    @staticmethod
    def save(
        output_dir: Path,
        provider: str,
        ats: ATSResult,
        profile: CVProfile,
        feedback: str,
        optimized_cv: str,
        processing_seconds: float,
    ) -> Path:
        output_dir.mkdir(parents=True, exist_ok=True)
        out_path = output_dir / "analysis.json"
        payload = {
            "timestamp_utc": datetime.now(timezone.utc).isoformat(),
            "provider": provider,
            "processing_seconds": round(processing_seconds, 2),
            "ats": asdict(ats),
            "cv_profile": asdict(profile),
            "feedback": feedback,
            "optimized_cv_text": optimized_cv,
        }
        out_path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
        return out_path


class OptimizedCVExporter:
    """Create optimized DOCX and PDF dynamically from rewritten CV text."""

    SECTION_NAMES: Final[set[str]] = {
        "professional summary", "summary", "profile", "technical skills", "skills", "core skills",
        "professional experience", "experience", "work experience", "employment history", "projects",
        "selected projects", "education", "certifications", "certification", "additional information",
        "achievements", "awards", "languages", "references",
    }

    @staticmethod
    def save_docx(optimized_cv_text: str, output_dir: Path) -> Path:
        output_dir.mkdir(parents=True, exist_ok=True)
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

        styles = doc.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(10)

        lines = [line.rstrip() for line in optimized_cv_text.splitlines()]
        non_empty_seen = 0

        for raw_line in lines:
            line = raw_line.strip()
            if not line:
                continue

            clean_heading = re.sub(r"[:\-]+$", "", line).strip().lower()
            is_bullet = bool(re.match(r"^(•|-|\*)\s+", line))
            is_heading = clean_heading in OptimizedCVExporter.SECTION_NAMES or (
                line.isupper() and len(line.split()) <= 5 and not is_bullet and non_empty_seen > 0
            )

            if non_empty_seen == 0:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(line)
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(16)
            elif non_empty_seen == 1 and ("@" in line or "|" in line or "+" in line or "linkedin" in line.lower()):
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(line)
                run.font.name = "Arial"
                run.font.size = Pt(9)
            elif is_heading:
                paragraph = doc.add_heading(re.sub(r"[:\-]+$", "", line).title(), level=2)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(11)
                    run.bold = True
            elif is_bullet:
                text = re.sub(r"^(•|-|\*)\s+", "", line).strip()
                paragraph = doc.add_paragraph(text, style="List Bullet")
                paragraph.paragraph_format.space_after = Pt(2)
            else:
                paragraph = doc.add_paragraph(line)
                paragraph.paragraph_format.space_after = Pt(4)

            non_empty_seen += 1

        out_path = output_dir / "optimized_cv.docx"
        doc.save(out_path)
        return out_path

    @staticmethod
    def save_pdf(optimized_cv_text: str, output_dir: Path) -> Path:
        """Create a simple ATS-friendly PDF from optimized text using reportlab."""
        try:
            from reportlab.lib.pagesizes import A4  # type: ignore
            from reportlab.lib.units import inch  # type: ignore
            from reportlab.pdfgen import canvas  # type: ignore
        except ImportError as exc:
            raise ImportError("Missing dependency: reportlab. Install with: pip install reportlab") from exc

        output_dir.mkdir(parents=True, exist_ok=True)
        pdf_path = output_dir / "optimized_cv.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=A4)
        width, height = A4
        margin = 0.6 * inch
        y = height - margin
        line_height = 12
        max_chars = 98

        lines = [line.rstrip() for line in optimized_cv_text.splitlines()]
        first_line = True
        for raw_line in lines:
            line = raw_line.strip()
            if not line:
                y -= line_height / 2
                continue

            is_bullet = bool(re.match(r"^(•|-|\*)\s+", line))
            clean_heading = re.sub(r"[:\-]+$", "", line).strip().lower()
            is_heading = clean_heading in OptimizedCVExporter.SECTION_NAMES or (line.isupper() and len(line.split()) <= 5 and not first_line)

            if first_line:
                c.setFont("Helvetica-Bold", 16)
                c.drawCentredString(width / 2, y, line[:80])
                y -= 18
                first_line = False
                continue

            if is_heading:
                c.setFont("Helvetica-Bold", 11)
                y -= 3
                c.drawString(margin, y, re.sub(r"[:\-]+$", "", line).title()[:90])
                y -= line_height
                continue

            c.setFont("Helvetica", 9)
            text = re.sub(r"^(•|-|\*)\s+", "- ", line) if is_bullet else line
            for wrapped in wrap(text, width=max_chars):
                if y < margin:
                    c.showPage()
                    y = height - margin
                    c.setFont("Helvetica", 9)
                c.drawString(margin, y, wrapped)
                y -= line_height

        c.save()
        return pdf_path


# =====================================================
# 8) MAIN APPLICATION
# =====================================================

@dataclass(frozen=True, slots=True)
class AppConfig:
    """Runtime configuration."""

    provider: str
    output_dir: Path
    export_docx: bool
    export_pdf: bool
    export_markdown: bool
    export_json: bool


class CVAnalyzerApp:
    """Backend application orchestration."""

    def __init__(self, config: AppConfig) -> None:
        self._config = config
        self._llm = LLMFactory.create(config.provider)
        self._scorer = ATSScorer()
        self._structure = CVStructureAnalyzer()
        self._feedback = FeedbackGenerator(self._llm)
        self._rewriter = CVRewriter(self._llm)

    def analyse(self, cv_path: str, jd_path: str) -> None:
        start = time.perf_counter()
        self._config.output_dir.mkdir(parents=True, exist_ok=True)
        log_path = setup_logging(self._config.output_dir)
        logging.info("Started CV analysis")
        logging.info("Provider requested: %s | LLM available: %s", self._config.provider, bool(self._llm))

        cv_text = read_document(cv_path)
        jd_text = read_document(jd_path)
        profile = self._structure.analyse(cv_text)
        ats = self._scorer.score(cv_text, jd_text, profile)

        prompt_1 = PromptBuilder.feedback_prompt(cv_text, jd_text, ats, profile)
        feedback = self._feedback.generate(prompt_1, ats, profile)

        # This is the key assignment requirement: use generated feedback to optimize the actual CV.
        prompt_2 = PromptBuilder.rewrite_prompt(cv_text, jd_text, feedback, ats)
        optimized_cv = self._rewriter.rewrite(prompt_2, cv_text, feedback, ats, profile)

        generated_files: list[Path] = []
        if self._config.export_markdown:
            md_path, txt_path = ReportExporter.save(feedback, self._config.output_dir)
            generated_files.extend([md_path, txt_path])

        if self._config.export_docx:
            docx_path = OptimizedCVExporter.save_docx(optimized_cv, self._config.output_dir)
            generated_files.append(docx_path)

        if self._config.export_pdf:
            pdf_path = OptimizedCVExporter.save_pdf(optimized_cv, self._config.output_dir)
            generated_files.append(pdf_path)

        elapsed = time.perf_counter() - start
        if self._config.export_json:
            json_path = JSONExporter.save(
                self._config.output_dir,
                self._config.provider if self._llm else "fallback-rule-based",
                ats,
                profile,
                feedback,
                optimized_cv,
                elapsed,
            )
            generated_files.append(json_path)

        logging.info("ATS overall score: %s", ats.overall_score)
        logging.info("Generated files: %s", [str(path) for path in generated_files])
        logging.info("Completed CV analysis in %.2f seconds", elapsed)

        print("\n=== AI CV Feedback Backend Result ===")
        print(f"Provider Used: {self._config.provider if self._llm else 'fallback-rule-based'}")
        print(f"Overall ATS Score: {ats.overall_score}%")
        print(f"Keyword Score: {ats.keyword_score}%")
        print(f"Skills Score: {ats.skills_score}%")
        print(f"Experience Score: {ats.experience_score}%")
        print(f"Formatting Score: {ats.formatting_score}%")
        print(f"Readability Score: {ats.readability_score}%")
        print(f"Section Score: {ats.section_score}%")
        print(f"Matched Keywords: {', '.join(ats.matched_keywords[:12]) or 'None'}")
        print(f"Missing Keywords: {', '.join(ats.missing_keywords[:12]) or 'None'}")
        print("\n--- Feedback Preview ---")
        print(feedback[:1500] + ("..." if len(feedback) > 1500 else ""))
        print("\n--- Files Generated ---")
        for path in generated_files:
            print(path)
        print(f"Log File: {log_path}")


# =====================================================
# 9) CLI
# =====================================================

def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="AI CV feedback and optimization backend")
    parser.add_argument("--cv", help="Path to CV file PDF/DOCX/TXT")
    parser.add_argument("--jd", help="Path to job description file PDF/DOCX/TXT")
    parser.add_argument("--provider", choices=["gemini", "openai"], default=DEFAULT_LLM_PROVIDER, help="LLM provider")
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT_DIR), help="Output folder path")
    parser.add_argument("--docx", action="store_true", help="Generate optimized CV DOCX")
    parser.add_argument("--pdf", action="store_true", help="Generate optimized CV PDF")
    parser.add_argument("--markdown", action="store_true", help="Generate feedback Markdown and TXT")
    parser.add_argument("--json", action="store_true", help="Generate analysis JSON")
    parser.add_argument("--no-export-cv", action="store_true", help="Do not generate optimized CV DOCX/PDF")
    return parser


def resolve_config(args: argparse.Namespace) -> AppConfig:
    """Create runtime config. Default behavior exports all required assignment files."""
    explicit_exports = any([args.docx, args.pdf, args.markdown, args.json])
    if args.no_export_cv:
        export_docx = False
        export_pdf = False
    else:
        export_docx = args.docx or not explicit_exports
        export_pdf = args.pdf or not explicit_exports

    return AppConfig(
        provider=args.provider,
        output_dir=Path(args.output),
        export_docx=export_docx,
        export_pdf=export_pdf,
        export_markdown=args.markdown or not explicit_exports,
        export_json=args.json or not explicit_exports,
    )


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()

    try:
        cv_path = args.cv or input("Enter CV file path (PDF/DOCX/TXT): ").strip()
        jd_path = args.jd or input("Enter Job Description file path (PDF/DOCX/TXT): ").strip()
        config = resolve_config(args)
        CVAnalyzerApp(config).analyse(cv_path, jd_path)
        return 0
    except (FileNotFoundError, ValueError, ImportError, OSError) as exc:
        print(f"ERROR: {exc}")
        logging.exception("Application error: %s", exc)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
